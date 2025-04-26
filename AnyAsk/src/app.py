#IMPORING NECESSARY LIBRARIES
import streamlit as st
from io import BytesIO
from docx import Document
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter

#CUSTOM CLASS
# from hugging_face_transformer_agent_calling import answer_question
from agent import a121_rag_chain
# from vector_store import create_vector_store
from vector_store_using_huggingface import create_vector_store

model_name = "HuggingFaceH4/zephyr-7b-beta"

# PROCESS_INPUT FUNCTION FIXES:
def process_input(input_type, input_data):
    # Processes different input types and returns a vectorstore.
    documents = ""
    
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        else:
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError("Expected string for 'Text' input type")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        else:
            doc = Document(BytesIO(input_data.read()))
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        else:
            text = str(input_data.read().decode('utf-8'))
        documents = text
    else:
        raise ValueError("Unsupported input type")
    
    #Split the document text into manageable chunks
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == 'Link':
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)
    
    return create_vector_store(texts)




def main():
    st.title('Advanced RAG Q&A System')
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])
    
    if input_type == "Link":
        number_input = st.number_input(min_value=1, max_value=20, step=1, label="Enter a number")
        input_data = []
        for i in range(number_input):
            url = st.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_area("Enter the text")  # Changed to text_area for multiline
    else:
        input_data = st.file_uploader(f"Upload a {input_type} file", type=[input_type.lower()])
    
    if st.button("Proceed") and input_data:
        try:
            vector_store = process_input(input_type, input_data)
            st.session_state["vector_store"] = vector_store  # Fixed variable name
            st.success("Documents processed successfully!")
        except Exception as e:
            st.error(f"Error processing documents: {str(e)}")
    
    if "vector_store" in st.session_state:
        query = st.text_input("Ask your question")
        if query and st.button("Submit"):
            answer = a121_rag_chain(st.session_state["vector_store"], query)
            st.write(answer['result'])  # Display just the answer
            with st.expander("See sources"):
                for doc in answer['source_documents']:
                    st.write(doc.page_content)
                    st.write("---")

    
if __name__=='__main__':
    main()